from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
import psycopg2
import urllib.request
import main

DB_INFO = {'dbname': 'postgres',
           'user': 'postgres',
           'password': '1',
           'host': 'localhost'
           }


def make_fio(first_name, second_name, last_name):
    """Создание Фамилия И.О."""
    return last_name.title() + " " + first_name[:1].title() + "." + \
           second_name[:1].title() + "."


def change_doc(type_doc, phone):
    """Изменение файла."""

    # Получение фамилии, имени, отчества
    print(phone)
    conn = psycopg2.connect(dbname=DB_INFO['dbname'],
                            user=DB_INFO['user'],
                            password=DB_INFO['password'],
                            host=DB_INFO['host'])
    cursor = conn.cursor()
    cursor.execute(f"""SELECT name, last_name, second_name 
                      FROM human 
                      WHERE number = '{phone}'""")
    res_select = cursor.fetchall()
    first_name = res_select[0][0]
    second_name = res_select[0][2]
    last_name = res_select[0][1]

    # Получение id типа документа и создание имени файла документа
    cursor.execute(f"""SELECT id_document 
                       FROM document
                       WHERE type = '{type_doc}'""")
    res_select = cursor.fetchall()
    id_type_doc = res_select[0][0]
    filename = "template_" + str(id_type_doc) + ".docx"
    cursor.close()
    conn.close() 

    # Запрос документа
    url_template = "http://redpix8i.beget.tech/templates/" + filename
    urllib.request.urlretrieve(url_template, "tmp.docx")
    # Имя фотки должно быть номер_телефона.png
    urllib.request.urlretrieve(f"http://redpix8i.beget.tech/images/{phone}.png", "tmp.png")

    fio = make_fio(first_name, second_name, last_name)
    doc = DocxTemplate("tmp.docx")
    context = {"fio": fio}
    doc.render(context)
    doc.save("tmp.docx")
    doc = Document("tmp.docx")
    p = doc.add_paragraph()
    r = p.add_run()
    r.font.name = "Times New Roman"
    r.font.bold = True
    r.add_text("Подпись: ")
    r.add_picture("tmp.png", width=Inches(1), height=Inches(1))
    # file_= {'file': ('tmp.doxc', open('tmp.dox', 'rb'))}
    # r = requests.post('http://redpix8i.beget.tech/documents/', files=file_)
    
    doc.save("tmp.docx")
    return 'tmp.docx'

